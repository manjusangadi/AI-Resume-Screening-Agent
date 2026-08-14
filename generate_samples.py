import os
import json

# Define job description
JOB_DESCRIPTION = """Title: Junior AI Research Associate & Python Engineer

Role Overview:
We are seeking a talented Junior AI Research Associate & Python Developer to join our team. The candidate will work on developing AI agents, fine-tuning machine learning models, and building clean data analysis pipelines.

Requirements:
- Education: Bachelor's, Master's, or PhD in Computer Science, Data Science, AI, or a related technical field.
- Experience: At least 2 years of professional software development or research experience.
- Technical Skills:
  - Strong proficiency in Python, Pandas, NumPy
  - Hands-on experience with Machine Learning (Scikit-Learn, PyTorch, or TensorFlow)
  - Practical understanding of NLP (TF-IDF, Cosine Similarity, embeddings, NLTK)
  - Database management using SQL
  - Experience building REST APIs (Flask or FastAPI)
- Soft Skills: Strong analytical thinking, problem-solving, and communication.
"""

# Define sample candidate profiles
CANDIDATES = [
    {
        "filename": "john_doe_outstanding.pdf",
        "name": "John Doe",
        "email": "john.doe@email.com",
        "phone": "+1-555-0199",
        "text": """JOHN DOE
Email: john.doe@email.com | Phone: +1-555-0199 | Address: New York, NY

OBJECTIVE
Highly motivated AI Engineer with 4 years of experience building Python-based machine learning pipelines, NLP similarity engines, and SQL data analytics solutions.

EDUCATION
- Master of Science in Data Science | Columbia University (2020 - 2022)
- Bachelor of Science in Computer Science | NYU (2016 - 2020)

EXPERIENCE
AI Engineer | TechSolutions Inc. (2022 - Present)
- Built an internal search platform using TF-IDF and Cosine Similarity to scan millions of documents.
- Developed machine learning models using PyTorch and Scikit-Learn.
- Wrote high-performance REST APIs in Flask and FastAPI to serve predictions.
- Managed database schemas and complex queries using PostgreSQL and SQL.

Python Developer | CoreCode Corp (2020 - 2022)
- Designed and maintained Python backend systems using Pandas and NumPy for high-throughput data analysis.
- Automated reporting pipelines that cut manual data prep time by 40%.

SKILLS
- Programming: Python, SQL, C++
- Libraries: Pandas, NumPy, Scikit-Learn, PyTorch, NLTK
- APIs & Databases: Flask, FastAPI, PostgreSQL, SQLite
- Concepts: TF-IDF, Cosine Similarity, Machine Learning, Natural Language Processing (NLP)
"""
    },
    {
        "filename": "sarah_jenkins_switcher.docx",
        "name": "Sarah Jenkins",
        "email": "sarah.jenkins@email.com",
        "phone": "+1-555-0188",
        "text": """SARAH JENKINS
Email: sarah.jenkins@email.com | Phone: +1-555-0188

SUMMARY
Experienced Web Developer transitioning into Data Science. Solid foundations in Python programming and SQL database design, currently building NLP projects and expanding machine learning knowledge.

EDUCATION
- Bachelor of Arts in Economics | Boston University (2018 - 2022)

EXPERIENCE
Frontend Developer | WebStyles Studio (2022 - Present)
- Built interactive dashboards in React and Node.js.
- Integrated REST APIs and worked closely with backend databases.
- Began integrating simple machine learning analytics into frontends using Python and Pandas.

Self-Employed / Freelancer | AI Learner (2022 - 2023)
- Created several open-source Python projects using TF-IDF and NLTK for text classification.
- Self-studied machine learning fundamentals using Scikit-Learn and NumPy.

SKILLS
- Languages: Python, JavaScript, HTML, CSS, SQL
- Data Science: Pandas, NumPy, TF-IDF, NLTK basics, Scikit-Learn
- Web Tech: React, Node.js, Express, Flask
"""
    },
    {
        "filename": "david_smith_junior.txt",
        "name": "David Smith",
        "email": "david.smith@email.com",
        "phone": "+1-555-0177",
        "text": """DAVID SMITH
Email: david.smith@email.com | Phone: +1-555-0177

CAREER SUMMARY
Junior Python developer eager to learn and grow. Passionate about software development, coding fundamentals, and building websites. Currently seeking a starting role in AI or Software Engineering.

EDUCATION
- Associate Degree in Information Technology | City College (2021 - 2023)

EXPERIENCE
Intern Developer | ByteCode Systems (6 months, 2023)
- Assisted with Python scripting and minor bug fixes.
- Wrote basic SQL queries for standard databases.
- Documented project setups and API routes.

SKILLS
- Programming: Python, HTML, CSS, SQL (Basic)
- Frameworks: Flask (Beginner)
- Tools: Git, VS Code, Windows OS
"""
    },
    {
        "filename": "emily_taylor_nlp.pdf",
        "name": "Emily Taylor",
        "email": "emily.taylor@email.com",
        "phone": "+1-555-0166",
        "text": """EMILY TAYLOR
Email: emily.taylor@email.com | Phone: +1-555-0166

NLP RESEARCH ASSOCIATE
AI researcher with 3 years of experience focusing on Natural Language Processing (NLP), text similarity, embeddings, and vector databases. Experienced in Python, PyTorch, and deploying machine learning pipelines.

EDUCATION
- Master of Science in Artificial Intelligence | Stanford University (2021 - 2023)
- Bachelor of Science in Mathematics | UCLA (2017 - 2021)

EXPERIENCE
NLP Engineer | SemanticAI Labs (2023 - Present)
- Implemented state-of-the-art NLP models using PyTorch, Hugging Face, and NLTK.
- Developed search algorithms leveraging Cosine Similarity and TF-IDF vectors.
- Automated data labeling pipelines using Pandas, reducing costs by 30%.

Python & Data Engineer | AnalyticsCorp (2021 - 2023)
- Managed complex analytical SQL databases and wrote custom Python scripts to ETL data.
- Built dashboard APIs with FastAPI and deployed dockerized apps.

SKILLS
- AI & NLP: Machine Learning, NLP, PyTorch, NLTK, TF-IDF, Cosine Similarity, Embeddings
- Programming: Python, R, SQL, Bash
- Software: Pandas, NumPy, Scikit-Learn, FastAPI, Flask, Docker, PostgreSQL
"""
    },
    {
        "filename": "michael_brown_web.txt",
        "name": "Michael Brown",
        "email": "michael.b@email.com",
        "phone": "+1-555-0155",
        "text": """MICHAEL BROWN
Email: michael.b@email.com | Phone: +1-555-0155

SENIOR WEB DEVELOPER
Full-stack web architect with 8 years of experience building secure, scalable web applications. Expert in Javascript frameworks and high-performance server architectures.

EDUCATION
- Bachelor of Science in Computer Science | University of Michigan (2012 - 2016)

EXPERIENCE
Lead Architect | WebScale Inc (2020 - Present)
- Designed full-stack platforms using React, Node.js, and MongoDB.
- Optimized database query structures resulting in a 50% increase in load times.
- Led a team of 5 junior and mid-level web developers.

Senior Backend Engineer | NetBuild LLC (2016 - 2020)
- Built RESTful APIs using Python, Django, and PostgreSQL.
- Handled CI/CD workflows and server configurations on AWS.

SKILLS
- Web Development: JavaScript, React, Node.js, Express, HTML, CSS, Tailwind CSS
- Databases: PostgreSQL, MongoDB, Redis
- Cloud & Tools: AWS, Docker, Git, Python, Django, REST APIs
"""
    },
    {
        "filename": "jessica_garcia_ds.pdf",
        "name": "Jessica Garcia",
        "email": "jessica.g@email.com",
        "phone": "+1-555-0144",
        "text": """JESSICA GARCIA
Email: jessica.g@email.com | Phone: +1-555-0144

DATA SCIENTIST
Statistical expert with 5 years of experience applying machine learning algorithms to corporate data challenges. Proficient in Python, SQL, statistical modeling, and data visualization.

EDUCATION
- PhD in Statistics | University of Washington (2016 - 2021)
- Bachelor of Science in Applied Math | UC Berkeley (2012 - 2016)

EXPERIENCE
Senior Data Scientist | FinMetrics Inc (2021 - Present)
- Created predictive machine learning models using Scikit-Learn, XGBoost, and Pandas.
- Wrote complex SQL scripts to extract features from multi-million row tables.
- Applied NLP techniques to analyze customer reviews and sentiments.

Data Analyst | BizData Solutions (2020 - 2021)
- Developed statistical dashboards and reports using Python, Pandas, and Tableau.

SKILLS
- Core: Python, SQL, R, SAS
- ML: Machine Learning, Scikit-Learn, TensorFlow, Pandas, NumPy, XGBoost
- Analytics: Statistics, NLP, Cosine Similarity, Data Visualization, Matplotlib, Tableau
"""
    },
    {
        "filename": "james_wilson_pm.docx",
        "name": "James Wilson",
        "email": "j.wilson@email.com",
        "phone": "+1-555-0133",
        "text": """JAMES WILSON
Email: j.wilson@email.com | Phone: +1-555-0133

TECHNICAL PRODUCT MANAGER
Product Manager with 6 years of experience driving development lifecycle of SaaS platforms. Strong background in user research, roadmapping, and agile management.

EDUCATION
- MBA | Harvard Business School (2018 - 2020)
- Bachelor of Science in Business Admin | Boston College (2012 - 2016)

EXPERIENCE
Product Manager | CloudWorks (2020 - Present)
- Managed the lifecycle of active cloud computing modules from concept to deployment.
- Coordinated cross-functional teams of engineers, design, and marketing.
- Defined product requirements, user stories, and acceptance criteria.

Project Manager | DevTech LLC (2016 - 2018)
- Managed schedule and budgets for corporate software deliveries.

SKILLS
- Product: Product Management, Agile, Scrum, Jira, Product Roadmap, Market Research
- Tech: SQL (Basic), HTML/CSS, Python (Conceptual)
"""
    },
    {
        "filename": "amanda_martinez_phd.pdf",
        "name": "Amanda Martinez",
        "email": "amanda.m@email.com",
        "phone": "+1-555-0122",
        "text": """AMANDA MARTINEZ
Email: amanda.m@email.com | Phone: +1-555-0122

AI RESEARCHER
Recent PhD graduate specializing in deep learning, neural networks, and advanced natural language processing. Deep theoretical understanding of transformers and text embeddings.

EDUCATION
- PhD in Computer Science (Focus on AI) | MIT (2020 - 2025)
- Bachelor of Science in Computer Science | UT Austin (2016 - 2020)

EXPERIENCE
Graduate Research Assistant | MIT AI Lab (2020 - 2025)
- Published 4 papers in top NLP/ML conferences on text similarity, TF-IDF enhancements, and Cosine Similarity limitations.
- Developed deep learning models in PyTorch and TensorFlow.
- Coded extensively in Python, leveraging NumPy, Pandas, and NLTK.

SKILLS
- Deep Tech: Machine Learning, Deep Learning, NLP, PyTorch, TensorFlow, Transformers
- Code: Python, C++, NumPy, Pandas, NLTK
- Math: Linear Algebra, Vector Embeddings, Cosine Similarity, Statistics
"""
    },
    {
        "filename": "robert_anderson_python.txt",
        "name": "Robert Anderson",
        "email": "robert.a@email.com",
        "phone": "+1-555-0111",
        "text": """ROBERT ANDERSON
Email: robert.a@email.com | Phone: +1-555-0111

PYTHON DEVELOPER
Software developer with 3 years of experience writing clean, maintainable code in Python. Experienced in building server apps, scraping web data, and API construction.

EDUCATION
- Bachelor of Science in Computer Science | Texas A&M (2017 - 2021)

EXPERIENCE
Python Developer | DevHouse Corp (2021 - Present)
- Built scraping engines and parsed HTML/XML databases with BeautifulSoup.
- Developed REST APIs in Flask.
- Maintained legacy script bases and migrated Python 2 scripts to Python 3.

SKILLS
- Languages: Python, Java, JavaScript
- Frameworks: Flask, Django
- Tools: BeautifulSoup, Git, PostgreSQL, Docker, Linux
"""
    },
    {
        "filename": "lisa_thomas_strong_no_degree.docx",
        "name": "Lisa Thomas",
        "email": "lisa.t@email.com",
        "phone": "+1-555-0100",
        "text": """LISA THOMAS
Email: lisa.t@email.com | Phone: +1-555-0100

SENIOR AI DEVELOPER
Self-taught AI enthusiast and developer with 5 years of experience building machine learning applications, NLP systems, and backend services. Proven track record of delivery without formal college degree.

EDUCATION
- Self-Directed Bootcamps & Certifications (Python, ML, NLP, SQL)

EXPERIENCE
AI Engineer | InnovateTech (2021 - Present)
- Developed ML classification engines using Scikit-Learn and Pandas.
- Implemented semantic document searches using NLTK, TF-IDF, and Cosine Similarity.
- Optimized SQLite/SQL server databases.
- Built APIs in Flask and FastAPI.

Backend Developer | StartupHub (2019 - 2021)
- Programmed Python microservices and managed ETL flows with Pandas and NumPy.

SKILLS
- AI & NLP: Machine Learning, NLP, TF-IDF, Cosine Similarity, Scikit-Learn, PyTorch, NLTK
- Programming: Python, SQL, Pandas, NumPy
- Web: Flask, FastAPI, REST APIs
"""
    },
    {
        "filename": "william_jackson_sales.pdf",
        "name": "William Jackson",
        "email": "william.j@email.com",
        "phone": "+1-555-0099",
        "text": """WILLIAM JACKSON
Email: william.j@email.com | Phone: +1-555-0099

SALES DIRECTOR
Results-driven sales leader with 10+ years of experience leading high-performance B2B sales teams in the technology sector. Proven history of exceeding sales quotas and expanding client portfolios.

EDUCATION
- Bachelor of Science in Business Management | Ohio State University (2010 - 2014)

EXPERIENCE
Regional Sales Director | CloudCore Inc (2019 - Present)
- Directed a team of 12 sales representatives, achieving 120% of quota.
- Developed sales pitches, client relationships, and business development plans.
- Handled negotiating contracts and pricing schemas.

Account Executive | SaaSFlow (2014 - 2019)
- Managed client onboarding and upsell pipelines.

SKILLS
- Core: Sales, Business Development, Negotiation, Lead Generation, Account Management
- CRM: Salesforce, Hubspot
- Tools: Excel, PowerPoint
"""
    }
]

def main():
    # Make directories
    os.makedirs("data/resumes", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    
    # Write job description
    with open("data/job_description.txt", "w", encoding="utf-8") as f:
        f.write(JOB_DESCRIPTION)
    print("Created data/job_description.txt")

    # Attempt imports for docx and reportlab to write true files
    has_docx = False
    try:
        import docx
        has_docx = True
    except ImportError:
        print("Warning: python-docx not installed. Will write plain text for docx files.")

    has_reportlab = False
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        has_reportlab = True
    except ImportError:
        print("Warning: reportlab not installed. Will write plain text for pdf files.")

    for cand in CANDIDATES:
        filename = cand["filename"]
        filepath = os.path.join("data/resumes", filename)
        
        # Determine active type
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".txt":
            # Write plain text
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(cand["text"])
            print(f"Created TXT resume: {filepath}")
            
        elif ext == ".docx":
            if has_docx:
                # Write true word document
                doc = docx.Document()
                doc.add_heading(cand["name"], level=0)
                
                # Split content into paragraphs
                lines = cand["text"].split("\n")
                for line in lines:
                    if line.strip():
                        if line.isupper() and len(line) < 30:
                            doc.add_heading(line, level=2)
                        else:
                            doc.add_paragraph(line)
                doc.save(filepath)
                print(f"Created DOCX resume: {filepath}")
            else:
                # Fallback to TXT with .txt extension
                fallback_path = os.path.splitext(filepath)[0] + ".txt"
                with open(fallback_path, "w", encoding="utf-8") as f:
                    f.write(cand["text"])
                print(f"Created DOCX fallback as TXT: {fallback_path}")
                
        elif ext == ".pdf":
            if has_reportlab:
                # Write true PDF document
                c = canvas.Canvas(filepath, pagesize=letter)
                c.setFont("Helvetica-Bold", 16)
                c.drawString(50, 750, cand["name"])
                
                c.setFont("Helvetica", 10)
                y = 730
                lines = cand["text"].split("\n")
                
                for line in lines:
                    if y < 50:
                        c.showPage()
                        c.setFont("Helvetica", 10)
                        y = 750
                    
                    line_str = line.strip()
                    if not line_str:
                        y -= 10
                        continue
                        
                    if line_str.isupper() and len(line_str) < 30:
                        c.setFont("Helvetica-Bold", 12)
                        y -= 15
                        c.drawString(50, y, line_str)
                        c.setFont("Helvetica", 10)
                        y -= 5
                    else:
                        # Wrap line if it's too long
                        if len(line_str) > 90:
                            parts = [line_str[i:i+90] for i in range(0, len(line_str), 90)]
                            for part in parts:
                                c.drawString(50, y, part)
                                y -= 12
                        else:
                            c.drawString(50, y, line_str)
                            y -= 12
                c.save()
                print(f"Created PDF resume: {filepath}")
            else:
                # Fallback to TXT with .txt extension
                fallback_path = os.path.splitext(filepath)[0] + ".txt"
                with open(fallback_path, "w", encoding="utf-8") as f:
                    f.write(cand["text"])
                print(f"Created PDF fallback as TXT: {fallback_path}")

    print("Sample generation complete. 11 resumes created successfully in data/resumes/")

if __name__ == "__main__":
    main()
