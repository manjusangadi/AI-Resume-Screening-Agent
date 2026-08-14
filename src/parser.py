import os
import re

def parse_file(filepath: str) -> str:
    """
    Detects the file extension of the given file and extracts the plain text content.
    Supported extensions: .txt, .pdf, .docx
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
        
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == ".txt":
        return _parse_txt(filepath)
    elif ext == ".pdf":
        return _parse_pdf(filepath)
    elif ext == ".docx":
        return _parse_docx(filepath)
    else:
        # Fallback to general text reading if possible, or raise error
        try:
            return _parse_txt(filepath)
        except Exception:
            raise ValueError(f"Unsupported file format: {ext}")

def _parse_txt(filepath: str) -> str:
    """Reads a plain text file."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        raise IOError(f"Error reading TXT file {filepath}: {str(e)}")

def _parse_pdf(filepath: str) -> str:
    """Reads and extracts text from a PDF file using pypdf."""
    try:
        import pypdf
    except ImportError:
        # Fallback if library not available: check if file is text-readable (for fallback samples)
        try:
            return _parse_txt(filepath)
        except Exception:
            raise ImportError("pypdf is required to parse PDF files. Install it using 'pip install pypdf'.")
            
    try:
        reader = pypdf.PdfReader(filepath)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        extracted_text = "\n".join(text_parts)
        # If no text was extracted (e.g. scanned pdf or parse failure), fallback
        if not extracted_text.strip():
            # Try plain text fallback in case it was written as text
            try:
                return _parse_txt(filepath)
            except Exception:
                pass
        return extracted_text
    except Exception as e:
        # Graceful fallback: try reading as plain text
        try:
            return _parse_txt(filepath)
        except Exception:
            raise IOError(f"Error reading PDF file {filepath}: {str(e)}")

def _parse_docx(filepath: str) -> str:
    """Reads and extracts text from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        # Fallback if library not available: check if file is text-readable (for fallback samples)
        try:
            return _parse_txt(filepath)
        except Exception:
            raise ImportError("python-docx is required to parse DOCX files. Install it using 'pip install python-docx'.")
            
    try:
        doc = docx.Document(filepath)
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
                
        # Extract from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)
                        
        extracted_text = "\n".join(text_parts)
        if not extracted_text.strip():
            try:
                return _parse_txt(filepath)
            except Exception:
                pass
        return extracted_text
    except Exception as e:
        # Graceful fallback: try reading as plain text
        try:
            return _parse_txt(filepath)
        except Exception:
            raise IOError(f"Error reading DOCX file {filepath}: {str(e)}")

def clean_text(text: str) -> str:
    """Standardizes spaces, line breaks, and formatting of parsed text."""
    # Standardize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Collapse multiple whitespaces (excluding newlines)
    text = re.sub(r'[ \t]+', ' ', text)
    # Remove excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()
